import os
import random
import sys
import win32com.client

import numpy as np
import docx
from docx.shared import Pt, Inches, Mm
import docx2pdf
from PIL import Image, ImageDraw, ImageFont
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

# from src.MxN import First
from src.shepley import Shepley as Sh
from src.Nash import Nash
from src.Nash import smesh_nesh
from src.MxN import Sedl
from src.MxN import NoSedl
from src.MxN import m2xn
from src.MxN import mx2
from src.MxN import m_n_matrix


def convertDocx2Pdf(groupCount:int):
    try:
        print(groupCount)

        # docx2pdf.convert(os.getcwd() + '\Task_' + str(self.group) + '.docx',
        #                  os.getcwd() + '\Task_' + str(self.group) + '.pdf')
        # docx2pdf.convert(os.getcwd() + '\Answers_' + str(self.group) + '.docx',
        #                  os.getcwd() + '\Answers_' + str(self.group) + '.pdf')
        # docx2pdf.convert(os.getcwd() + '\Response_' + str(self.group) + '.docx',
        #                  os.getcwd() + '\Response_' + str(self.group) + '.pdf')
        # os.remove(os.path.abspath('Response_' + str(self.group) + '.docx'))
        # os.remove(os.path.abspath('Task_' + str(self.group) + '.docx'))
        # os.remove(os.path.abspath('Answers_' + str(self.group) + '.docx'))
        for count in range (groupCount):
            print(count)
            wdFormatPDF = 17

            inputFile = os.path.abspath('Task_' + str(count+1) + '.docx')
            outputFile = os.path.abspath('Task_' + str(count+1) + '.pdf')
            file = open(outputFile, "w")
            file.close()
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(inputFile)
            doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
            doc.Close()
            os.remove(inputFile)

            inputFile = os.path.abspath('Answers_' + str(count+1) + '.docx')
            outputFile = os.path.abspath('Answers_' + str(count+1) + '.pdf')
            file = open(outputFile, "w")
            file.close()
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(inputFile)
            doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
            doc.Close()
            os.remove(inputFile)

            inputFile = os.path.abspath('Response_' + str(count+1) + '.docx')
            outputFile = os.path.abspath('Response_' + str(count+1) + '.pdf')
            file = open(outputFile, "w")
            file.close()
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(inputFile)
            doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()
            os.remove(inputFile)

        #
        # os.remove(os.path.abspath('Task_' + str(self.group) + '.docx'))
        # os.remove(os.path.abspath('Answers_' + str(self.group) + '.docx'))
        return "Успешная генерация"
    except:
        return "Ошибка генерации"


class CreatePDF:

    def __init__(self, count: int, students: list, group: int):
        self.count = count
        self.students = students
        self.docTask = docx.Document()
        self.docAnswers = docx.Document()
        self.docResponse = docx.Document()
        self.group = group

    def title(self):
        try:
            sys.stderr = open("consoleoutput.log", "w")
            style = self.docTask.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)
            section = self.docTask.sections[0]
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(5)

            section = self.docAnswers.sections[0]
            section.right_margin = Mm(10)
            section.left_margin = Mm(10)
            section.top_margin = Mm(16.5)
            section.orientation = WD_ORIENT.PORTRAIT
            # print(section.orientation)

            # 5 Блоков в один блок записываются ответы даже на пары заданий с указанием из номера
            tableAnswer = self.docAnswers.add_table(rows=self.count + 1, cols=5, style='Table Grid')
            self.initAsnswersTable(tableAnswer, self.count)
            for counter in range(self.count):
                title = self.docTask.add_table(rows=1, cols=3)
                cell = title.cell(0, 0)
                cell.text = "Группа: " + str(self.group)
                cell.width = Inches(1)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                cell = title.cell(0, 1)
                cell.text = "ФИО: " + self.students[counter]
                cell.width = Inches(8)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                cell = title.cell(0, 2)
                cell.text = "№Вар. " + str(counter + 1)
                cell.width = Inches(1)
                cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                # title = self.docTask.add_paragraph()
                # run = title.add_run('Группа: 0 \t ФИО: '+ self.students[counter] + "\t\t\t№Вар.  25").font.size = Pt(10)
                # # r_fmt = run.font
                # # # run.font.size = Pt(10)
                # title.alignment = 3

                rand = random.choice([1,2,3])
                print(rand)

                self.taskOne(tableAnswer, counter,rand)
                self.taskTwo(tableAnswer, counter)
                self.taskThree(tableAnswer, counter)
                self.taskFour(tableAnswer, counter)
                # self.docTask.add_page_break()
                self.responsePage(counter)
                # Завершающая часть программы
                if counter + 1 != self.count:
                    self.docTask.add_page_break()
                else:
                    # Вместо self.count пойдет название группы
                    self.docTask.save("Task_" + str(self.group) + ".docx")
                    self.docAnswers.save("Answers_" + str(self.group) + ".docx")
                    self.docResponse.save("Response_" + str(self.group) + ".docx")
                    return 'Файлы сохранены в формате "docx"'

                    # try:
                    #
                    #     # docx2pdf.convert(os.getcwd() + '\Task_' + str(self.group) + '.docx',
                    #     #                  os.getcwd() + '\Task_' + str(self.group) + '.pdf')
                    #     # docx2pdf.convert(os.getcwd() + '\Answers_' + str(self.group) + '.docx',
                    #     #                  os.getcwd() + '\Answers_' + str(self.group) + '.pdf')
                    #     # docx2pdf.convert(os.getcwd() + '\Response_' + str(self.group) + '.docx',
                    #     #                  os.getcwd() + '\Response_' + str(self.group) + '.pdf')
                    #     # os.remove(os.path.abspath('Response_' + str(self.group) + '.docx'))
                    #     # os.remove(os.path.abspath('Task_' + str(self.group) + '.docx'))
                    #     # os.remove(os.path.abspath('Answers_' + str(self.group) + '.docx'))
                    #     wdFormatPDF = 17
                    #
                    #     inputFile = os.path.abspath('Task_' + str(self.group) + '.docx')
                    #     outputFile = os.path.abspath('Task_' + str(self.group) + '.pdf')
                    #     file = open(outputFile, "w")
                    #     file.close()
                    #     word = win32com.client.Dispatch("Word.Application")
                    #     doc = word.Documents.Open(inputFile)
                    #     doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
                    #     doc.Close()
                    #     os.remove(inputFile)
                    #
                    #     inputFile = os.path.abspath('Answers_' + str(self.group) + '.docx')
                    #     outputFile = os.path.abspath('Answers_' + str(self.group) + '.pdf')
                    #     file = open(outputFile, "w")
                    #     file.close()
                    #     word = win32com.client.Dispatch("Word.Application")
                    #     doc = word.Documents.Open(inputFile)
                    #     doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
                    #     doc.Close()
                    #     os.remove(inputFile)
                    #
                    #     inputFile = os.path.abspath('Response_' + str(self.group) + '.docx')
                    #     outputFile = os.path.abspath('Response_' + str(self.group) + '.pdf')
                    #     file = open(outputFile, "w")
                    #     file.close()
                    #     word = win32com.client.Dispatch("Word.Application")
                    #     doc = word.Documents.Open(inputFile)
                    #     doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
                    #     doc.Close()
                    #     word.Quit()
                    #     os.remove(inputFile)
                    #
                    #     #
                    #     # os.remove(os.path.abspath('Task_' + str(self.group) + '.docx'))
                    #     # os.remove(os.path.abspath('Answers_' + str(self.group) + '.docx'))
                    #     return "Успешная генерация"
                    # except:
                    #     return "Ошибка генерации"
        except:
            print("Ошибка генерации")

    def initAsnswersTable(self, tableAnswer, count: int):
        style = self.docAnswers.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)

        tableAnswer.cell(0, 1).text = "Часть 1"
        tableAnswer.cell(0, 2).text = "Часть 2"
        tableAnswer.cell(0, 3).text = "Часть 3"
        tableAnswer.cell(0, 4).text = "Часть 4"

        for i in range(count):
            cell = tableAnswer.cell(i + 1, 0)
            cell.text = str(i + 1) + ". " + self.students[i]

    def taskOne(self, tableAnswer, counter: int,rand:int):
        try:

            taskOne = self.docTask.add_paragraph('Часть 1. Матричные игры с нулевой суммой 2х2\n')
            taskOne.add_run('Задание: \n'
                            '1. Для заданных двух матриц найти нижнюю и верхнюю цену игры. \n'
                            '2. В каждой из рассмотренных игр установить возможность решений в чистых стратегиях. \n'
                            '3. Найти оптимальные решения в каждой игре расчетным путем.').font.size = Pt(9)
            gameOne = np.random.randint(0, 50, (2, 2))
            gameTwo = np.random.randint(0, 50, (2, 2))
            if rand == 1:
                playOne = Sedl.Matrix_2x2(gameOne)
                playOne.sed_point()
                playTwo = NoSedl.Matrix_2x2(gameTwo)
                playTwo.sed_point()
                tableAnswer.cell(counter + 1,
                                 1).text = f"Игра №1:\nSa= ({playOne.S1[0]}, {playOne.S1[1]}) \n Sb=({playOne.S2[0]}, {playOne.S2[1]}),\n" \
                                           f" Седловая точка=({playOne.dot})\n" \
                                           f"V= {playOne.A}\n" \
                                           f"Игра №2:\n Sa= ({playTwo.Sa[0]}, {playTwo.Sa[1]}) \n Sb=({playTwo.Sb[0]}," \
                                           f" {playTwo.Sb[1]}),\n V= {playTwo.v} "
                tableForTask(playOne.array, playTwo.array, 2, 2)
            elif rand == 2:
                playOne = NoSedl.Matrix_2x2(gameOne)
                playOne.sed_point()
                playTwo = NoSedl.Matrix_2x2(gameTwo)
                playTwo.sed_point()
                tableAnswer.cell(counter + 1,
                                 1).text = f"Игра №1:\n Sa= ({playOne.Sa[0]}, {playOne.Sa[1]}) \n Sb=({playOne.Sb[0]}," \
                                           f" {playOne.Sb[1]}),\n V= {playOne.v}\n" \
                                           f"Игра №2:\n Sa= ({playTwo.Sa[0]}, {playTwo.Sa[1]}) \n Sb=({playTwo.Sb[0]}," \
                                           f" {playTwo.Sb[1]}),\n V= {playTwo.v} "
            elif rand == 3:
                playTwo = Sedl.Matrix_2x2(gameOne)
                playTwo.sed_point()
                playOne = NoSedl.Matrix_2x2(gameTwo)
                playOne.sed_point()
                tableAnswer.cell(counter + 1,
                                 1).text = f"Игра №1:\nSa= ({playOne.Sa[0]}, {playOne.Sa[1]}) \n Sb=({playOne.Sb[0]}," \
                                           f" {playOne.Sb[1]}),\n V= {playOne.v}\n" \
                                           f"Игра №2:\nSa= ({playTwo.S1[0]}, {playTwo.S1[1]}) \n Sb=({playTwo.S2[0]}," \
                                           f" {playTwo.S2[1]}),\n" \
                                           f" Седловая точка=({playTwo.dot}) "

            tableForTask(playOne.array, playTwo.array, 2, 2)

            paragraph = self.docTask.add_paragraph()
            run = paragraph.add_run("Игра №1                                                 Игра №2\n")
            run.font.size = Pt(9)
            run.add_picture('table_1.png')
            run_2 = paragraph.add_run("\t \t \t")
            run_2.add_picture('table_2.png')

            # self.docTask.save("task.docx")
            # self.docAnswers.save("Answers.docx")

            os.remove(os.path.abspath('table_1.png'))
            os.remove(os.path.abspath('table_2.png'))
            print("Успешная генерация части 1")

        except:
            print('Ошибка генерации части 1')

    def taskTwo(self, tableAnswer, counter):
        try:
            taskTwo = self.docTask.add_paragraph('Часть 2. Матричные игры с нулевой суммой 2хn, mx2 и mxn\n')
            taskTwo.add_run('Задание: \n'
                            '1. Для заданных двух матриц найти нижнюю и верхнюю цену игры. \n'
                            '2. В каждой из рассмотренных игр установить возможность решений в чистых стратегиях. \n'
                            '3. Найти оптимальные решения в каждой игре расчетным путем.').font.size = Pt(9)
            gameOne = m2xn.initFunc()
            gameTwo = mx2.initFunc()
            matrix = m_n_matrix.init(4, 6)
            # print(matrix)
            gameThree = m_n_matrix.Matrix_mxn(matrix)
            gameThree.execute()
            # print(gameThree.matrix)
            # print(gameOne.array[0])
            # print(gameTwo.array)

            tableForTask(gameOne.array, gameTwo.array, 2, 4)
            tableForTaskOnMxN(gameThree.matrix)
            # print(gameThree.p[4])
            tableAnswer.cell(counter + 1,
                             2).text = f"Игра 2xn:\n Sa= ({round(gameOne.S1[0], 3)}, {round(gameOne.S1[1], 3)}) \n" \
                                       f" Sb=({round(gameOne.S2[0], 3)}, {round(gameOne.S2[1], 3)}," \
                                       f"{round(gameOne.S2[2], 3)}, {round(gameOne.S2[3], 3)}),\n" \
                                       f" V=({round(gameOne.V, 3)})\n" \
                                       f"Игра mx2:\n Sa= ({round(gameTwo.S1[0], 3)}, {round(gameTwo.S1[1], 3)}," \
                                       f"{round(gameTwo.S1[2], 3)}, {round(gameTwo.S1[3], 3)}) \n Sb=({round(gameTwo.S2[0], 4)}" \
                                       f" {round(gameTwo.S2[1], 3)}),\n V=({round(gameTwo.V, 3)})\n" \
                                       f"Игра mxn:\n" \
                                       f"Sa= ({gameThree.p[0]}, {gameThree.p[1]}, " \
                                       f"{gameThree.p[2]},  {gameThree.p[3]}) \n Sb=({gameThree.q[0]}, " \
                                       f"{gameThree.q[1]},  {gameThree.q[2]},  {gameThree.q[3]},  {gameThree.q[4]}," \
                                       f" {gameThree.q[5]}),\n Va=({gameThree.va}),\n Vb=({gameThree.vb}) "

            paragraph = self.docTask.add_paragraph()
            run = paragraph.add_run("Игра 2xn                     "
                                    "                           Игра mx2      "
                                    "          Игра mxn\n")
            run.font.size = Pt(9)
            run.add_picture('table_1.png')
            run_2 = paragraph.add_run("\t ")
            run_2.add_picture('table_2.png')
            run_3 = paragraph.add_run("\t ")
            run_3.add_picture('table_3.png')

            os.remove(os.path.abspath('table_1.png'))
            os.remove(os.path.abspath('table_2.png'))
            os.remove(os.path.abspath('table_3.png'))
            print("Успешная генерация части 2")
        except:
            print("Ошибка генерации части 2")

    def taskThree(self, tableAnswer, counter):
        try:
            taskThree = self.docTask.add_paragraph('Часть 3. Биматричные одношаговые игры 2x2.'
                                                   'Равновесие по Нэшу в чистых и смешанных стратегиях\n')
            taskThree.add_run('Задание: \n'
                              '1. Для игры №1, найти графическим путем Парето-множество решений, отметив здесь же решения, '
                              'соответствующие равновесию по Нэшу. \n'
                              '2. Для игры №2, убедитесь в отсутствии оптимальных по Нэшу решений в чистых стратегиях отметив здесь же решения, '
                              'соответствующие равновесию по Нэшу. \n'
                              '3. Найти оптимальное по Нэшу решение в смешанных стартегиях. \n'
                              '4. Определить средний выйгрыш 1-го и 2-го игроков при равновесии по Нэшу ').font.size = Pt(
                9)

            playOne = Nash.Nash()
            playOne.generate_game()
            playTwo = smesh_nesh.initFunction()
            playTwo.execute()
            tableForTaskThree(playOne.A, playOne.B, playTwo.A, playTwo.B)

            paragraph = self.docTask.add_paragraph()
            run = paragraph.add_run("Игра №1                         "
                                    "                                        Игра №2\n")
            run.font.size = Pt(9)
            run.add_picture('table_1.png')
            run_2 = paragraph.add_run("\t \t \t")
            run_2.add_picture('table_2.png')

            os.remove(os.path.abspath('table_1.png'))
            os.remove(os.path.abspath('table_2.png'))
            if not playOne.Nash:
                tableAnswer.cell(counter + 1, 3).text = "Игра №1:\nРавновесий по Нэшу нет\n" + "Оптимум по Парето: "
                for obj in playOne.pareto[:-1]:
                    tableAnswer.cell(counter + 1, 3).text += str(obj) + ","
                tableAnswer.cell(counter + 1, 3).text += str(playOne.pareto[-1])
                # tableAnswer.cell(counter+1,3).width = Inches(0.5)
            else:
                tableAnswer.cell(counter + 1, 3).text = "Игра №1: \nРавновесие по Нэшу: ("
                for obj in playOne.Nash[:-1]:
                    tableAnswer.cell(counter + 1, 3).text += str(obj) + ","
                tableAnswer.cell(counter + 1, 3).text += playOne.Nash[-1] + ")"
                for obj in playOne.answer[:-1]:
                    tableAnswer.cell(counter + 1, 3).text += str(obj) + ","
                tableAnswer.cell(counter + 1, 3).text += playOne.answer[-1]
                tableAnswer.cell(counter + 1, 3).text += "\nОптимум по Парето: ("
                for obj in playOne.pareto[:-1]:
                    tableAnswer.cell(counter + 1, 3).text += str(obj) + ","
                tableAnswer.cell(counter + 1, 3).text += playOne.pareto[-1] + ")"

            tableAnswer.cell(counter + 1, 3).text += "\nИгра №2:\n"
            tableAnswer.cell(counter + 1, 3).text += f"Sa =({round(playTwo.SA[0], 3)}, {round(playTwo.SA[1], 3)}) \n" \
                                                     f"Sb=({round(playTwo.SB[0], 3)}, {round(playTwo.SB[1], 3)})\n" \
                                                     f"V = ({round(playTwo.VA, 3)}, {round(playTwo.VB, 3)})\n"
            tableAnswer.cell(counter + 1, 3).text += "Оптимум по Парето: (\n"

            for obj in playTwo.pareto[:-1]:
                tableAnswer.cell(counter + 1, 3).text += str(obj) + ","
            tableAnswer.cell(counter + 1, 3).text += playTwo.pareto[-1] + ")"

            print("Успешная генерация части 3")
        except:
            print("Ошибка генерации части 3")


    def taskFour(self, tableAnswer, counter):

        try:
            taskFive = self.docTask.add_paragraph(
                'Часть 4. Дележ в кооперативных играх (вектор Шепли, С - ядро) \n')
            taskFive.add_run('Задание: \n'
                             '1. Определить выигрыши каждого из игроков в случае их объединения на основе'
                             ' использования вектора Шепли.\n'
                             '2. Проверить принадлежность вектора Шепли С - ядру ').font.size = Pt(9)

            function = {}
            vGame = {0: "V(1)", 1: "V(2)", 2: "(V3)", 3: "V(1,2)", 4: "V(1,3)", 5: "V(2,3)", 6: "V(1,2,3)"}
            for j in range(1, 4):
                function[str(j)] = random.randint(100, 400)
            function['12'] = function['1'] + function['2'] + random.randint(20, 200)
            function['13'] = function['1'] + function['3'] + random.randint(20, 200)
            function['23'] = function['2'] + function['3'] + random.randint(20, 200)
            function['123'] = function['1'] + function['2'] + function['3'] + random.randint(20, 200)
            # print(f"Входные данные: {function}")
            vec = Sh.Shepley_vector(function, 3)
            vec.find_vecotr()


            tableAnswer.cell(counter + 1,
                             4).text = f"q1= ({round(vec.vector_find[0], 3)}),\n q2=({round(vec.vector_find[1], 3)}),\n" \
                                       f" q3=({round(vec.vector_find[2], 3)}),\n ({vec.core})"

            tableForTaskSheply(function, vGame)
            paragraph = self.docTask.add_paragraph()
            run = paragraph.add_run("\n")
            run.font.size = Pt(9)
            run.add_picture('table.png')
            os.remove(os.path.abspath('table.png'))
            print("Успешная генерация Части 4")
        except:
            print("Ошибка генерации Части 4")

    def responsePage(self, counter):
        try:

            style = self.docResponse.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            section = self.docResponse.sections[0]
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(5)
            section.right_margin = Mm(10)
            section.left_margin = Mm(10)

            table = self.docResponse.add_table(rows=1, cols=3)

            cell = table.cell(0, 0)
            cell.text = "Группа: " + str(self.group)
            cell.width = Inches(1)
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            cell.paragraphs[0].runs[0].font.size = Pt(10)

            cell = table.cell(0, 1)
            cell.text = "ФИО: " + self.students[counter]
            cell.width = Inches(8)
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)

            cell = table.cell(0, 2)
            cell.text = "№Вар. " + str(counter + 1)
            cell.width = Inches(1)
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            cell.paragraphs[0].runs[0].font.size = Pt(10)

            self.docResponse.add_paragraph('\nЧасть 1. Матричные игры с нулевой суммой 2х2\n\n') \
                .add_run(
                'Игра №1\n1. SA = (                          ),\t SB = (                       ),\t цена игры =  \n\n'
                'Игра №2\n2. SA = (                          ),\t SB = (                       ),\t цена игры = \n') \
                .font.size = Pt(13)

            self.docResponse.add_paragraph('Часть 2. Матричные игры с нулевой суммой 2хn, mx2 и mxn\n\n') \
                .add_run(
                'Игра mx2\n1. SA = (                                          ), SB = (                                          ), цена игры =  \n'
                '\nИгра 2xn\n2. SA = (                                          ), SB = (                                          ), цена игры =\n'
                '\nИгра mхn\n3. SA = (                                          ), SB = (                                          ), '
                'цена игры = \n') \
                .font.size = Pt(13)

            self.docResponse.add_paragraph(
                'Часть 3. Биматричные одношаговые игры 2x2.Равновесие по Нэшу в чистых и смещанных стратегиях\n\n') \
                .add_run(
                'Игра №1\n1. Равновесие по Нэшу  ______________________________\n'
                '2. Оптимально по Парето ______________________________\n'
                '\nИгра №2\n1. Равновесие по Нэшу  ______________________________\n'
                '2. Оптимально по Парето ______________________________') \
                .font.size = Pt(13)

            self.docResponse.add_paragraph(
                'Часть 4. Дележ в кооперативных играх (вектор Шепли, С - ядро) \n\n') \
                .add_run(
                'Вектор Шепли') \
                .font.size = Pt(13)
            self.docResponse.add_table(rows=1, cols=3, style="Table Grid")
            self.docResponse.add_paragraph('\n') \
                .add_run('Принадлежность С – ядру ___________________________________________')

            if counter + 1 != self.count:
                self.docResponse.add_page_break()
            print("Генерация бланка ответов прошла успешно")
        except:
            print("Не Работает генерация бланка ответов")


def tableForTaskOnMxN(playThree: list):
    try:
        tables = docx.Document()
        table3 = tables.add_table(rows=4, cols=6, style='Table Grid')
        for row in range(4):
            for col in range(6):
                cell = table3.cell(row, col)
                cell.text = str(playThree[row][col])
                cell.width = Inches(0.2)
        createPngOnMxN(table3, 30, 18)
        print("Успешная генерация таблицы MxN")
    except:
        print("Ошибка генерации таблицы MxN")


# Создает таблицы для первой группы задач для их последующего конверта в png
def tableForTask(playOne: list, playTwo: list, rows: int, collums: int):
    try:
        tables = docx.Document()
        table1 = tables.add_table(rows=rows, cols=collums, style='Table Grid')

        tables.add_paragraph(' ')
        if rows == 2 and collums == 2:
            table2 = tables.add_table(rows=rows, cols=collums, style='Table Grid')
        else:
            table2 = tables.add_table(rows=collums, cols=rows, style='Table Grid')
        if rows == 2 and collums == 2:
            for row in range(rows):
                for col in range(collums):
                    cell = table1.cell(row, col)
                    cell.text = str(playOne[row][col])
                    cell.width = Inches(0.2)

            for row in range(rows):
                for col in range(collums):
                    cell = table2.cell(row, col)
                    cell.text = str(playTwo[row][col])
                    cell.width = Inches(0.2)
        else:
            for row in range(rows):
                for col in range(collums):
                    cell = table1.cell(row, col)
                    cell.text = str(playOne[row][col])
                    cell.width = Inches(0.2)

            for row in range(collums):
                for col in range(rows):
                    cell = table2.cell(row, col)
                    cell.text = str(playTwo[row][col])
                    cell.width = Inches(0.2)

        createPng(table1, table2, 30, 18)
        print('Успешная генерация таблиц для части 1')
    except:
        print('Ошибка генерации таблиц для части 1')


def tableForTaskThree(playOneA: list, playOneB: list, playTwoA: list, playTwoB: list):
    tables = docx.Document()
    table1 = tables.add_table(rows=2, cols=2, style='Table Grid')
    table2 = tables.add_table(rows=2, cols=2, style='Table Grid')
    rowA = table1.rows[0]
    rowB = table1.rows[1]
    rowA.cells[0].text = '(' + str(playOneA[0][0]) + ' ; ' + str(playOneB[0][0]) + ')'
    rowA.cells[1].text = '(' + str(playOneA[0][1]) + ' ; ' + str(playOneB[0][1]) + ')'
    rowB.cells[0].text = '(' + str(playOneA[1][0]) + ' ; ' + str(playOneB[1][0]) + ')'
    rowB.cells[1].text = '(' + str(playOneA[1][1]) + ' ; ' + str(playOneB[1][1]) + ')'

    rowA = table2.rows[0]
    rowB = table2.rows[1]
    rowA.cells[0].text = '(' + str(playTwoA[0][0]) + ' ; ' + str(playTwoB[0][0]) + ')'
    rowA.cells[1].text = '(' + str(playTwoA[0][1]) + ' ; ' + str(playTwoB[0][1]) + ')'
    rowB.cells[0].text = '(' + str(playTwoA[1][0]) + ' ; ' + str(playTwoB[1][0]) + ')'
    rowB.cells[1].text = '(' + str(playTwoA[1][1]) + ' ; ' + str(playTwoB[1][1]) + ')'

    createPng(table1, table2, 50, 20)


def tableForTaskSheply(play: list, game: dict):
    try:
        tables = docx.Document()
        table = tables.add_table(rows=2, cols=7, style='Table Grid')
        for col in range(7):
            cell = table.cell(0, col)
            cell.text = game.get(col)
            cell.width = Inches(0.5)

        row = table.rows[1]
        row.cells[0].text = str(play['1'])
        row.cells[1].text = str(play['2'])
        row.cells[2].text = str(play['3'])
        row.cells[3].text = str(play['12'])
        row.cells[4].text = str(play['13'])
        row.cells[5].text = str(play['23'])
        row.cells[6].text = str(play['123'])
        createPngSheply(table, 45, 20)
        print("Успешная генерация таблицы для части 4")
    except:
        print("Ошибка генерации таблицы для части 4")


def createPngSheply(table, cell_width, cell_height):
    num_rows = len(table.rows)
    num_cols = len(table.columns)
    # Вычисляем размер изображения на основе размеров таблицы и ячеек

    img_width = num_cols * cell_width
    img_height = num_rows * cell_height
    # Создаем изображение с учетом размеров таблицы
    img = Image.new('RGB', (img_width + 1, img_height + 1), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype('arial.ttf', 10)

    # Рисуем содержимое таблицы на изображении
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Рисуем ячейку таблицы
            draw.rectangle([(j * cell_width, i * cell_height), ((j + 1) * cell_width, (i + 1) * cell_height)],
                           outline=(0, 0, 0))
            draw.text((j * cell_width + 5, i * cell_height + 5), cell.text, fill=(0, 0, 0), font=font)

    # Сохраняем изображение в формате png
    img.save('table.png')


def createPngOnMxN(table3, cell_width, cell_height):
    num_rows = len(table3.rows)
    num_cols = len(table3.columns)
    # Вычисляем размер изображения на основе размеров таблицы и ячеек

    img_width = num_cols * cell_width
    img_height = num_rows * cell_height
    # Создаем изображение с учетом размеров таблицы
    img = Image.new('RGB', (img_width + 1, img_height + 1), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype('arial.ttf', 10)

    # Рисуем содержимое таблицы на изображении
    for i, row in enumerate(table3.rows):
        for j, cell in enumerate(row.cells):
            # Рисуем ячейку таблицы
            draw.rectangle([(j * cell_width, i * cell_height), ((j + 1) * cell_width, (i + 1) * cell_height)],
                           outline=(0, 0, 0))
            draw.text((j * cell_width + 5, i * cell_height + 5), cell.text, fill=(0, 0, 0), font=font)

    # Сохраняем изображение в формате png
    img.save('table_3.png')


# Создает png Таблицы для ее дальнейшей вставки в файл
def createPng(table1, table2, cell_width, cell_height):
    try:
        # Определяем размеры таблицы
        num_rows = len(table1.rows)
        num_cols = len(table1.columns)
        # Вычисляем размер изображения на основе размеров таблицы и ячеек

        img_width = num_cols * cell_width
        img_height = num_rows * cell_height
        # Создаем изображение с учетом размеров таблицы
        img = Image.new('RGB', (img_width + 1, img_height + 1), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype('arial.ttf', 10)

        # Рисуем содержимое таблицы на изображении
        for i, row in enumerate(table1.rows):
            for j, cell in enumerate(row.cells):
                # Рисуем ячейку таблицы
                draw.rectangle([(j * cell_width, i * cell_height), ((j + 1) * cell_width, (i + 1) * cell_height)],
                               outline=(0, 0, 0))
                draw.text((j * cell_width + 5, i * cell_height + 5), cell.text, fill=(0, 0, 0), font=font)

        # Сохраняем изображение в формате png
        img.save('table_1.png')

        num_rows = len(table2.rows)
        num_cols = len(table2.columns)
        # Вычисляем размер изображения на основе размеров таблицы и ячеек

        img_width = num_cols * cell_width
        img_height = num_rows * cell_height
        # Создаем изображение с учетом размеров таблицы
        img = Image.new('RGB', (img_width + 1, img_height + 1), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype('arial.ttf', 10)

        # Рисуем содержимое таблицы на изображении
        for i, row in enumerate(table2.rows):
            for j, cell in enumerate(row.cells):
                # Рисуем ячейку таблицы
                draw.rectangle([(j * cell_width, i * cell_height), ((j + 1) * cell_width, (i + 1) * cell_height)],
                               outline=(0, 0, 0))
                draw.text((j * cell_width + 5, i * cell_height + 5), cell.text, fill=(0, 0, 0), font=font)

        # Сохраняем изображение в формате png
        img.save('table_2.png')
        print('Успешная генепация картинок части 1')
    except:
        print("Ошибка генерации картинок части 1")




